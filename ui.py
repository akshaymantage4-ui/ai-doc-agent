import streamlit as st
import tempfile

from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain_core.documents import Document

from docx import Document as DocxDocument
from PIL import Image
import pytesseract


# ⚠️ Set Tesseract path (change if different)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# -------------------------
# Load different file types
# -------------------------
def load_document(file):
    file_type = file.name.split(".")[-1].lower()

    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(file.read())
        file_path = tmp.name

    # PDF
    if file_type == "pdf":
        loader = PyMuPDFLoader(file_path)
        docs = loader.load()
        for d in docs:
            d.metadata["source"] = file.name
        return docs

    # DOCX
    elif file_type == "docx":
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": file.name})]

    # TXT
    elif file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file.name})]

    # IMAGE (OCR)
    elif file_type in ["jpg", "jpeg", "png"]:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return [Document(page_content=text, metadata={"source": file.name})]

    return []


# -------------------------
# Split text
# -------------------------
def split_docs(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )
    return splitter.split_documents(documents)


# -------------------------
# Embeddings
# -------------------------
def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )


# -------------------------
# Vector DB (cached)
# -------------------------
@st.cache_resource
def process_docs(documents):
    chunks = split_docs(documents)
    embeddings = get_embeddings()
    return FAISS.from_documents(chunks, embeddings)


# -------------------------
# Search
# -------------------------
def search_docs(query, vectorstore):
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    return retriever.invoke(query)


# -------------------------
# LLM
# -------------------------
def get_llm():
    return ChatGroq(
        model="llama-3.1-8b-instant",
        api_key="# 🔴 PUT YOUR KEY HERE"   
    )


# -------------------------
# Ask Question
# -------------------------
def ask_question(query, docs, llm):
    context = "\n".join([doc.page_content for doc in docs])

    prompt = f"""
    Answer the question based only on the context below.

    Context:
    {context}

    Question:
    {query}
    """

    response = llm.invoke(prompt)

    sources = [doc.metadata.get("source", "Unknown") for doc in docs]

    return f"{response.content}\n\n📄 Sources: {', '.join(set(sources))}"


# -------------------------
# Summary
# -------------------------
def summarize_docs(docs, llm):
    context = "\n".join([doc.page_content for doc in docs[:5]])

    prompt = f"Summarize this document:\n{context}"
    response = llm.invoke(prompt)

    return response.content


# -------------------------
# Insights
# -------------------------
def extract_insights(docs, llm):
    context = "\n".join([doc.page_content for doc in docs[:5]])

    prompt = f"""
    Extract key insights, important points, and conclusions:

    {context}
    """

    response = llm.invoke(prompt)
    return response.content


# -------------------------
# UI SETTINGS
# -------------------------
st.set_page_config(page_title="AI Document Assistant", layout="wide")

st.title("📄 AI Document Assistant")

with st.sidebar:
    st.title("⚙️ Options")
    st.write("Upload files and interact with AI")


# -------------------------
# Chat memory
# -------------------------
if "messages" not in st.session_state:
    st.session_state.messages = []


# -------------------------
# Upload files
# -------------------------
uploaded_files = st.file_uploader(
    "Upload Documents",
    type=["pdf", "docx", "txt", "jpg", "png"],
    accept_multiple_files=True
)

all_docs = []

if uploaded_files:
    st.write("### Uploaded Files:")
    for file in uploaded_files:
        st.write(f"📄 {file.name}")
        docs = load_document(file)
        all_docs.extend(docs)

    vectorstore = process_docs(all_docs)
    st.success("Documents processed successfully!")


# -------------------------
# Buttons
# -------------------------
if uploaded_files:

    llm = get_llm()

    if st.button("📄 Summarize"):
        summary = summarize_docs(all_docs, llm)
        st.write("### Summary:")
        st.write(summary)

    if st.button("🔍 Extract Insights"):
        insights = extract_insights(all_docs, llm)
        st.write("### Insights:")
        st.write(insights)


# -------------------------
# Chat UI
# -------------------------
query = st.chat_input("Ask something about your documents...")

if query and uploaded_files:
    st.session_state.messages.append({"role": "user", "content": query})

    llm = get_llm()
    relevant_docs = search_docs(query, vectorstore)
    answer = ask_question(query, relevant_docs, llm)

    st.session_state.messages.append({"role": "assistant", "content": answer})


# -------------------------
# Display chat
# -------------------------
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])
